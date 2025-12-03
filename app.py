from flask import Flask, render_template, request, send_file, flash, redirect, url_for
import os
import zipfile
from werkzeug.utils import secure_filename
from docx import Document
import io
import xml.etree.ElementTree as ET
from datetime import datetime
import base64

# PDF işleme için
try:
    import fitz  # PyMuPDF
    from PIL import Image as PILImage
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False

# ReportLab için (UDF to PDF)
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
    REPORTLAB_SUPPORT = True
except ImportError:
    REPORTLAB_SUPPORT = False

app = Flask(__name__)
app.config['SECRET_KEY'] = 'udf-converter-secret-key-2024'
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['OUTPUT_FOLDER'] = 'outputs'
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50MB max

ALLOWED_EXTENSIONS = {'docx', 'udf', 'pdf'}

os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['OUTPUT_FOLDER'], exist_ok=True)

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

UDF_TEMPLATE = '''<?xml version="1.0" encoding="UTF-8" ?>
<template format_id="1.8">
<content><![CDATA[{content}]]></content>
<properties><pageFormat mediaSizeName="1" leftMargin="42.51968479156494" rightMargin="28.34645652770996" topMargin="14.17322826385498" bottomMargin="14.17322826385498" paperOrientation="1" headerFOffset="20.0" footerFOffset="20.0" /></properties>
<elements resolver="hvl-default">
{elements}
</elements>
<styles><style name="default" description="Geçerli" family="Dialog" size="12" bold="false" italic="false" foreground="-13421773" FONT_ATTRIBUTE_KEY="javax.swing.plaf.FontUIResource[family=Dialog,name=Dialog,style=plain,size=12]" /><style name="hvl-default" family="Times New Roman" size="12" description="Gövde" /></styles>
</template>'''

def docx_to_udf_converter(docx_path, udf_path):
    """DOCX to UDF converter"""
    try:
        document = Document(docx_path)
        content = []
        elements = []
        current_offset = 0
        
        for paragraph in document.paragraphs:
            para_text = paragraph.text
            if not para_text:
                para_text = '\u200B'
            
            alignment = "0"
            if paragraph.alignment:
                if paragraph.alignment == 1:
                    alignment = "1"
                elif paragraph.alignment == 2:
                    alignment = "2"
                elif paragraph.alignment == 3:
                    alignment = "3"
            
            content.append(para_text)
            para_element = f'<paragraph Alignment="{alignment}" LeftIndent="0.0" RightIndent="0.0">'
            
            for run in paragraph.runs:
                run_text = run.text
                if run_text:
                    font_size = "12"
                    if run.font.size:
                        font_size = str(int(run.font.size.pt))
                    
                    font_family = run.font.name if run.font.name else "Times New Roman"
                    style_attrs = [
                        f'family="{font_family}"',
                        f'size="{font_size}"',
                        f'startOffset="{current_offset}"',
                        f'length="{len(run_text)}"'
                    ]
                    
                    if run.bold:
                        style_attrs.append('bold="true"')
                    if run.italic:
                        style_attrs.append('italic="true"')
                    
                    para_element += f'<content {" ".join(style_attrs)} />'
                    current_offset += len(run_text)
            
            if not paragraph.runs:
                para_element += f'<content startOffset="{current_offset}" length="{len(para_text)}" family="Times New Roman" size="12" />'
                current_offset += len(para_text)
            
            para_element += '</paragraph>'
            elements.append(para_element)
        
        if not content:
            content.append('\u200B')
            elements.append('<paragraph Alignment="0" LeftIndent="0.0" RightIndent="0.0"><content startOffset="0" length="1" family="Times New Roman" size="12" /></paragraph>')
        
        udf_content = UDF_TEMPLATE.format(
            content=''.join(content),
            elements='\n'.join(elements)
        )
        
        with zipfile.ZipFile(udf_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
            zipf.writestr('content.xml', udf_content)
        
        return True, "✅ DOCX başarıyla UDF'ye dönüştürüldü!"
    except Exception as e:
        return False, f"❌ Hata: {str(e)}"

def udf_to_docx_converter(udf_path, docx_path):
    """UDF to DOCX converter"""
    try:
        with zipfile.ZipFile(udf_path, 'r') as z:
            if 'content.xml' not in z.namelist():
                return False, "❌ Geçersiz UDF dosyası"
            
            with z.open('content.xml') as f:
                tree = ET.parse(f)
                root = tree.getroot()
        
        content_elem = root.find('content')
        if content_elem is None or content_elem.text is None:
            return False, "❌ İçerik bulunamadı"
        
        content_text = content_elem.text
        if content_text.startswith('<![CDATA[') and content_text.endswith(']]>'):
            content_text = content_text[9:-3]
        
        document = Document()
        elements_elem = root.find('elements')
        
        if elements_elem is not None:
            for para_elem in elements_elem.findall('paragraph'):
                paragraph = document.add_paragraph()
                alignment = para_elem.get('Alignment', '0')
                
                if alignment == '1':
                    paragraph.alignment = 1
                elif alignment == '2':
                    paragraph.alignment = 2
                elif alignment == '3':
                    paragraph.alignment = 3
                
                for content_elem in para_elem.findall('content'):
                    start_offset = int(content_elem.get('startOffset', '0'))
                    length = int(content_elem.get('length', '0'))
                    text = content_text[start_offset:start_offset+length]
                    run = paragraph.add_run(text)
                    
                    if content_elem.get('bold') == 'true':
                        run.bold = True
                    if content_elem.get('italic') == 'true':
                        run.italic = True
                    
                    font_size = content_elem.get('size')
                    if font_size:
                        from docx.shared import Pt
                        run.font.size = Pt(float(font_size))
        else:
            document.add_paragraph(content_text)
        
        document.save(docx_path)
        return True, "✅ UDF başarıyla DOCX'e dönüştürüldü!"
    except Exception as e:
        return False, f"❌ Hata: {str(e)}"

def pdf_to_udf_converter(pdf_path, udf_path):
    """PDF to UDF converter"""
    if not PDF_SUPPORT:
        return False, "❌ PDF desteği yüklü değil"
    
    try:
        pdf_document = fitz.open(pdf_path)
        content = []
        elements = []
        current_offset = 0

        for page_num in range(len(pdf_document)):
            page = pdf_document[page_num]
            text = page.get_text()
            
            if text:
                content.append(text)
                elements.append(f'<paragraph Alignment="0" LeftIndent="0.0" RightIndent="0.0"><content startOffset="{current_offset}" length="{len(text)}" family="Times New Roman" size="12" /></paragraph>')
                current_offset += len(text)
            
            image_list = page.get_images(full=True)
            for img_index, img in enumerate(image_list):
                xref = img[0]
                base_image = pdf_document.extract_image(xref)
                image_bytes = base_image["image"]
                
                image = PILImage.open(io.BytesIO(image_bytes))
                buffered = io.BytesIO()
                image.save(buffered, format="PNG")
                img_str = base64.b64encode(buffered.getvalue()).decode()
                
                placeholder = '\uFFFC'
                content.append(placeholder)
                elements.append(f'<image family="Times New Roman" size="10" imageData="{img_str}" startOffset="{current_offset}" length="1" />')
                current_offset += 1
            
            if page_num < len(pdf_document) - 1:
                content.append('\n')
                elements.append(f'<paragraph Alignment="0" LeftIndent="0.0" RightIndent="0.0"><content startOffset="{current_offset}" length="1" /></paragraph>')
                current_offset += 1

        udf_content = UDF_TEMPLATE.format(
            content=''.join(content),
            elements='\n'.join(elements)
        )

        with zipfile.ZipFile(udf_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
            zipf.writestr('content.xml', udf_content)
        
        return True, f"✅ PDF başarıyla UDF'ye dönüştürüldü! ({len(pdf_document)} sayfa)"
    except Exception as e:
        return False, f"❌ Hata: {str(e)}"

def udf_to_pdf_converter(udf_path, pdf_path):
    """UDF to PDF converter"""
    if not REPORTLAB_SUPPORT:
        return False, "❌ PDF oluşturma desteği yüklü değil"
    
    try:
        with zipfile.ZipFile(udf_path, 'r') as z:
            if 'content.xml' not in z.namelist():
                return False, "❌ Geçersiz UDF dosyası"
            
            with z.open('content.xml') as f:
                tree = ET.parse(f)
                root = tree.getroot()
        
        content_elem = root.find('content')
        if content_elem is None or content_elem.text is None:
            return False, "❌ İçerik bulunamadı"
        
        content_text = content_elem.text
        if content_text.startswith('<![CDATA[') and content_text.endswith(']]>'):
            content_text = content_text[9:-3]
        
        pdf = SimpleDocTemplate(pdf_path, pagesize=A4)
        styles = getSampleStyleSheet()
        pdf_elements = []
        
        elements_elem = root.find('elements')
        if elements_elem is not None:
            for para_elem in elements_elem.findall('paragraph'):
                alignment_map = {'0': TA_LEFT, '1': TA_CENTER, '2': TA_RIGHT, '3': TA_JUSTIFY}
                alignment = alignment_map.get(para_elem.get('Alignment', '0'), TA_LEFT)
                
                para_style = ParagraphStyle('CustomStyle', parent=styles['Normal'], alignment=alignment)
                para_text = ""
                
                for content_elem in para_elem.findall('content'):
                    start_offset = int(content_elem.get('startOffset', '0'))
                    length = int(content_elem.get('length', '0'))
                    text = content_text[start_offset:start_offset+length]
                    
                    if content_elem.get('bold') == 'true' and content_elem.get('italic') == 'true':
                        text = f"<b><i>{text}</i></b>"
                    elif content_elem.get('bold') == 'true':
                        text = f"<b>{text}</b>"
                    elif content_elem.get('italic') == 'true':
                        text = f"<i>{text}</i>"
                    
                    para_text += text
                
                if para_text.strip():
                    pdf_elements.append(Paragraph(para_text, para_style))
                    pdf_elements.append(Spacer(1, 5))
        
        pdf.build(pdf_elements)
        return True, "✅ UDF başarıyla PDF'e dönüştürüldü!"
    except Exception as e:
        return False, f"❌ Hata: {str(e)}"

def process_batch_conversion(files, conversion_type):
    """Process multiple files"""
    results = []
    output_files = []
    
    for file in files:
        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(filepath)
            
            base_name = os.path.splitext(filename)[0]
            
            if conversion_type == 'docx_to_udf':
                output_filename = f"{base_name}.udf"
                output_path = os.path.join(app.config['OUTPUT_FOLDER'], output_filename)
                success, message = docx_to_udf_converter(filepath, output_path)
            elif conversion_type == 'udf_to_docx':
                output_filename = f"{base_name}.docx"
                output_path = os.path.join(app.config['OUTPUT_FOLDER'], output_filename)
                success, message = udf_to_docx_converter(filepath, output_path)
            elif conversion_type == 'pdf_to_udf':
                output_filename = f"{base_name}.udf"
                output_path = os.path.join(app.config['OUTPUT_FOLDER'], output_filename)
                success, message = pdf_to_udf_converter(filepath, output_path)
            elif conversion_type == 'udf_to_pdf':
                output_filename = f"{base_name}.pdf"
                output_path = os.path.join(app.config['OUTPUT_FOLDER'], output_filename)
                success, message = udf_to_pdf_converter(filepath, output_path)
            else:
                success = False
                message = "Geçersiz dönüştürme tipi"
                output_path = None
            
            results.append({
                'filename': filename,
                'success': success,
                'message': message,
                'output': output_filename if success else None
            })
            
            if success:
                output_files.append(output_path)
            
            os.remove(filepath)
    
    return results, output_files

@app.route('/')
def index():
    return render_template('index.html', pdf_support=PDF_SUPPORT, reportlab_support=REPORTLAB_SUPPORT)

@app.route('/convert', methods=['POST'])
def convert():
    is_batch = request.form.get('is_batch') == 'true'
    
    if is_batch:
        files = request.files.getlist('files')
        conversion_type = request.form.get('conversion_type')
        
        if not files or len(files) == 0:
            flash('Dosya seçilmedi', 'error')
            return redirect(url_for('index'))
        
        results, output_files = process_batch_conversion(files, conversion_type)
        
        if output_files:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            zip_filename = f'converted_files_{timestamp}.zip'
            zip_path = os.path.join(app.config['OUTPUT_FOLDER'], zip_filename)
            
            with zipfile.ZipFile(zip_path, 'w') as zipf:
                for file_path in output_files:
                    zipf.write(file_path, os.path.basename(file_path))
                    os.remove(file_path)
            
            return send_file(zip_path, as_attachment=True, download_name=zip_filename)
        else:
            flash('Hiçbir dosya dönüştürülemedi', 'error')
            return redirect(url_for('index'))
    else:
        if 'file' not in request.files:
            flash('Dosya yüklenmedi', 'error')
            return redirect(url_for('index'))
        
        file = request.files['file']
        conversion_type = request.form.get('conversion_type')
        
        if file.filename == '':
            flash('Dosya seçilmedi', 'error')
            return redirect(url_for('index'))
        
        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(filepath)
            
            base_name = os.path.splitext(filename)[0]
            
            if conversion_type == 'docx_to_udf':
                output_filename = f"{base_name}.udf"
                output_path = os.path.join(app.config['OUTPUT_FOLDER'], output_filename)
                success, message = docx_to_udf_converter(filepath, output_path)
            elif conversion_type == 'udf_to_docx':
                output_filename = f"{base_name}.docx"
                output_path = os.path.join(app.config['OUTPUT_FOLDER'], output_filename)
                success, message = udf_to_docx_converter(filepath, output_path)
            elif conversion_type == 'pdf_to_udf':
                output_filename = f"{base_name}.udf"
                output_path = os.path.join(app.config['OUTPUT_FOLDER'], output_filename)
                success, message = pdf_to_udf_converter(filepath, output_path)
            elif conversion_type == 'udf_to_pdf':
                output_filename = f"{base_name}.pdf"
                output_path = os.path.join(app.config['OUTPUT_FOLDER'], output_filename)
                success, message = udf_to_pdf_converter(filepath, output_path)
            else:
                flash('Geçersiz dönüştürme tipi', 'error')
                return redirect(url_for('index'))
            
            os.remove(filepath)
            
            if success:
                return send_file(output_path, as_attachment=True, download_name=output_filename)
            else:
                flash(message, 'error')
                return redirect(url_for('index'))
        else:
            flash('Geçersiz dosya türü', 'error')
            return redirect(url_for('index'))

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)